import os
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from PIL import Image
import math

from log import log
from doc import add_table, set_font


def one_of_one(options):
  '''
  一格單張的函數
  '''
  doc = Document()
  doc = set_font(doc)
  for root, dirs, files in os.walk('Image'):
    # 返回地址、裡面所有資料夾、裡面所有檔案，故第二項用不到
    log().info(f'讀取到下列檔案：{files}')
    no = 0  # 序號

    for file in files:
      log().debug(f'讀取檔案：{file}')
      if file.lower().endswith(('.png', '.jpg', '.jpeg')):
        # 判斷副檔名是否為圖片檔
        no += 1  # 序號加一
        img_path = os.path.join(root, file)
        img = Image.open(img_path)
        show_type = 'H'  # 預設設定為高9公分
        if img.width / img.height > 15 / 9:
          # 大於預設寬高比15：9，設定寬為15公分。
          show_type = 'W'
        doc = add_table(doc, options, no, img_path, show_type=show_type)

      else:
        log().info(f'{file}並非支援的圖片檔')
  save(doc, options)


def lot_of_one(options):
  '''
  一格多張的函數
  '''
  doc = Document()
  doc = set_font(doc)
  for root, dirs, files in os.walk('Image'):
    log().info(f'讀取到下列檔案：{files}')
    no = 0  # 序號
    IMG_LIST = []  # 圖片的暫存清單，最多三個
    for file in files:
      log().debug(f'讀取檔案：{file}')
      if len(IMG_LIST) == 3:
        # 如果發現暫存清單已滿，先匯出
        no += 1  # 序號加一
        doc = add_table(doc, options, no, img_path=IMG_LIST, mode=2)
        IMG_LIST = []  # 清空暫存清單

      if file.lower().endswith(('.png', '.jpg', '.jpeg')):
        img_path = os.path.join(root, file)
        img = Image.open(img_path)

        if img.width / img.height <= 9 / 16:
          # 小於9/16，代表為手機圖片，儲存進暫存清單
          IMG_LIST.append(img_path)
        else:
          # 如果非手機圖片，則判斷是否有暫存清單，有則先匯出
          if IMG_LIST:
            no += 1  # 序號加一
            doc = add_table(doc, options, no, img_path=IMG_LIST, mode=2)
            IMG_LIST = []  # 清空暫存清單

          # 排除暫存清單後，依一般圖片處理
          no += 1
          show_type = 'H'  # 預設設定為高9公分
          if img.width / img.height > 15 / 9:
            # 大於預設寬高比15：9，設定寬為15公分。
            show_type = 'W'
          doc = add_table(doc, options, no, img_path, show_type=show_type)

      else:
        log().info(f'{file}並非支援的圖片檔')

    if IMG_LIST:
      # 如果最後清單不為空，則最後再儲存一次
      no += 1  # 序號加一
      doc = add_table(doc, options, no, img_path=IMG_LIST, mode=2)
      IMG_LIST = []  # 清空暫存清單
  save(doc, options)


def save(doc,options):
  header = doc.sections[0].header
  title = header.paragraphs[0].add_run(options['title'])
  title.font.size = Pt(20)
  title.font.bold = True
  header.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
  try:
    doc.save(f'{options["title"]}.docx')
    log().info('執行完畢')
    log().info(f'檔案已儲存為「{options["title"]}.docx」')
  except:
    log().critical('儲存失敗，請檢查有無相同檔名的檔案未關閉！', exc_info=True)


def crop_img():
  for root, dirs, files in os.walk('Image'):
    # 返回地址、裡面所有資料夾、裡面所有檔案，故第二項用不到
    log().debug(f'讀取到下列檔案：{files}')
    for file in files:
      log().debug(f'讀取檔案：{file}')
      if file.lower().endswith(('.png', '.jpg', '.jpeg')):
        # 判斷副檔名是否為圖片檔
        img_path = os.path.join(root, file)
        img = Image.open(img_path)  # 打開圖片
        w = img.width
        h = img.height
        if w / h <= 9 / 21:
          log().debug(f'{file}尺寸為{img.size}')
          # 如果圖片寬高比小於手機螢幕尺寸，判斷是長截圖
          nh = w * 2  # 以寬為基礎，每2倍的寬分割一次，
          blocks = math.ceil(h / nh)  # 應分割區塊數
          for i in range(0, blocks):
            cropped = img.crop((0, i * nh, w, (i+1) * nh))  # (原點x, 原點y, 終點x, 終點-y)
            cropped.save(os.path.join(root, f'{file}_block_{i}.jpg'))
          log().info(f'{file}共分割成{blocks}個檔案')
          os.remove(img_path)
          log().debug(f'{file}已刪除')

