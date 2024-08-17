import os
from docx import Document
from docx.shared import Inches, Cm
from PIL import Image

from log import log




# # 调整照片大小的函数
# def resize_image(image_path, target_width):
#   '''
#   :param image_path:圖片路徑
#   :param target_width:目標寬度
#   :return:返回調整後的圖片
#   '''
#   img = Image.open(image_path)
#   width_percent = (target_width / float(img.size[0]))
#   height_size = int((float(img.size[1]) * float(width_percent)))
#   resized_img = img.resize((target_width, height_size), Image.LANCZOS)
#   return resized_img
#
#
# # 存储照片的文件夹路径
# photos_folder = "Image"  # 将 "path_to_photos_folder" 替换为你存储照片的文件夹路径
#
# # 创建一个新的Word文档
# doc = Document()
#
# # 依次处理照片
# photo_count = 0
# photos_per_page = 3
# page_width = 12  # Word页面宽度，单位为Inches
#
# for root, _, files in os.walk(photos_folder):
#   for filename in files:
#     if filename.lower().endswith(('.png', '.jpg', '.jpeg')):
#       photo_count += 1
#       photo_path = os.path.join(root, filename)
#
#       # 调整照片大小
#       target_width = int(page_width / photos_per_page * 1440)  # 1440是Word中1英寸对应的像素数
#       resized_img = resize_image(photo_path, target_width)
#
#       # 将照片添加到Word文档中
#       if photo_count % photos_per_page == 1:
#         # 添加新页面
#         doc.add_page_break()
#
#       # 添加照片到当前页面
#       # doc.add_picture(photo_path, width=Inches(page_width / photos_per_page))
#       doc.add_picture(photo_path, height=Cm(9.5))
#       # print(photo_path)
#       # doc.add_picture('img.png', width=Inches(1440))
#
#       # 保存调整大小后的照片
#       # resized_photo_path = os.path.splitext(photo_path)[0] + "_resized.jpg"
#       # resized_img.save(resized_photo_path)
#
# # 保存Word文档
# doc.save("output.docx")
